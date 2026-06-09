"""ATS-safe .docx CV generation using python-docx.

Produces a single-column document with standard section headings, no tables,
text boxes, or columns — the layout that modern ATS parsers handle most
reliably. This is the primary output format for tailored CVs.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from cv_builder.models import UserProfile
from cv_builder.utils import build_filename, format_date_range

_RIGHT_TAB = Inches(6.5)
_DARK = RGBColor(0x22, 0x22, 0x22)
_GREY = RGBColor(0x55, 0x55, 0x55)


class DocxGenerator:
    """Render a :class:`UserProfile` as an ATS-friendly .docx CV."""

    def __init__(self, profile: UserProfile):
        self.profile = profile

    def generate(self, output_dir: str = ".", filename: str | None = None) -> str:
        pi = self.profile.personal_info
        if not pi:
            raise ValueError("Personal information is required to generate a CV")

        os.makedirs(output_dir, exist_ok=True)
        filename = filename or build_filename(pi.full_name, "docx")
        filepath = os.path.join(output_dir, filename)

        doc = Document()
        self._configure(doc)
        self._header(doc)
        self._summary(doc)
        self._experience(doc)
        self._education(doc)
        self._skills(doc)
        self._projects(doc)
        self._certifications(doc)
        self._additional(doc)
        doc.save(filepath)
        return filepath

    # ─── Setup ──────────────────────────────────────────────────────────────

    def _configure(self, doc: Document) -> None:
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

    # ─── Sections ───────────────────────────────────────────────────────────

    def _header(self, doc: Document) -> None:
        pi = self.profile.personal_info
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_p.add_run(pi.full_name)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = _DARK
        name_p.paragraph_format.space_after = Pt(2)

        contact = [pi.email, pi.phone, pi.location]
        if pi.linkedin:
            contact.append(pi.linkedin)
        if pi.portfolio:
            contact.append(pi.portfolio)
        contact = [c for c in contact if c]
        c_p = doc.add_paragraph()
        c_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_run = c_p.add_run("  |  ".join(contact))
        c_run.font.size = Pt(9.5)
        c_run.font.color.rgb = _GREY
        c_p.paragraph_format.space_after = Pt(6)

    def _heading(self, doc: Document, title: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = _DARK

    def _entry(self, doc: Document, title: str, org: str, location: str,
               start: str, end: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(_RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
        title_run = p.add_run(title)
        title_run.bold = True
        if start:
            p.add_run("\t" + format_date_range(start, end))

        if org:
            sub = doc.add_paragraph()
            sub.paragraph_format.space_after = Pt(2)
            sub_text = org if not location else f"{org}  ·  {location}"
            sub_run = sub.add_run(sub_text)
            sub_run.italic = True
            sub_run.font.size = Pt(10)
            sub_run.font.color.rgb = _GREY

    def _bullet(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(text)

    def _summary(self, doc: Document) -> None:
        pi = self.profile.personal_info
        if not pi or not pi.summary:
            return
        self._heading(doc, "Professional Summary")
        doc.add_paragraph(pi.summary)

    def _experience(self, doc: Document) -> None:
        if not self.profile.work_experience:
            return
        self._heading(doc, "Work Experience")
        for exp in self.profile.work_experience:
            self._entry(doc, exp.job_title, exp.company, exp.location,
                        exp.start_date, exp.end_date)
            for ach in exp.achievements:
                parts = [ach.description]
                if ach.quantified_impact:
                    parts.append(f"Impact: {ach.quantified_impact}")
                if ach.tools_used:
                    parts.append(f"Tools: {ach.tools_used}")
                self._bullet(doc, ". ".join(parts))

    def _education(self, doc: Document) -> None:
        if not self.profile.education:
            return
        self._heading(doc, "Education")
        for edu in self.profile.education:
            self._entry(doc, f"{edu.degree} in {edu.field_of_study}",
                        edu.institution, "", edu.start_date, edu.end_date)
            extras = []
            if edu.gpa:
                extras.append(f"GPA: {edu.gpa}")
            if edu.honors:
                extras.append(edu.honors)
            if edu.notable_coursework:
                extras.append(edu.notable_coursework)
            if extras:
                doc.add_paragraph("  •  ".join(extras))

    def _skills(self, doc: Document) -> None:
        sk = self.profile.skills
        if not (sk.technical or sk.soft or sk.languages):
            return
        self._heading(doc, "Skills")
        if sk.technical:
            self._labeled(doc, "Technical", ", ".join(s.name for s in sk.technical))
        if sk.soft:
            self._labeled(doc, "Soft Skills", ", ".join(sk.soft))
        if sk.languages:
            self._labeled(
                doc,
                "Languages",
                ", ".join(f"{language.name} ({language.proficiency})" for language in sk.languages),
            )

    def _projects(self, doc: Document) -> None:
        if not self.profile.projects:
            return
        self._heading(doc, "Projects")
        for proj in self.profile.projects:
            self._entry(doc, proj.name, proj.role, "", "", "")
            desc = proj.description
            if proj.outcomes:
                desc += f" {proj.outcomes}"
            self._bullet(doc, desc)
            if proj.technologies:
                self._labeled(doc, "Tech", proj.technologies)
            if proj.link:
                self._labeled(doc, "Link", proj.link)

    def _certifications(self, doc: Document) -> None:
        certs = self.profile.skills.certifications
        if not certs:
            return
        self._heading(doc, "Certifications")
        for cert in certs:
            self._bullet(doc, f"{cert.name} — {cert.issuing_org} ({cert.date})")

    def _additional(self, doc: Document) -> None:
        ai = self.profile.additional_info
        if not (ai.volunteer_work or ai.publications or ai.interests):
            return
        self._heading(doc, "Additional Information")
        if ai.volunteer_work:
            self._labeled(doc, "Volunteer", ai.volunteer_work)
        if ai.publications:
            self._labeled(doc, "Publications", ai.publications)
        if ai.interests:
            self._labeled(doc, "Interests", ai.interests)

    def _labeled(self, doc: Document, label: str, value: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)
